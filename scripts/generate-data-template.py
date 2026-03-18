#!/usr/bin/env python3
"""HACCP-ONE 실 운영 데이터 입력 템플릿 워드 문서 생성"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ── 헬퍼 함수 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E7D32')
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = '맑은 고딕'
            if ri % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_json_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # 배경색
    shading = p._element.get_or_add_pPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F0',
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x00)
    run.font.name = '맑은 고딕'


# ============================================================================
# 문서 시작
# ============================================================================

# 표지
title = doc.add_heading('HACCP-ONE 실 운영 데이터 입력 템플릿', level=0)
for run in title.runs:
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
run = p.add_run('이 문서의 각 테이블 양식에 맞춰 데이터를 작성해 주세요.\n입력 순서: 거래처 → 원재료 → 제품 → 레시피 → 매입(입고) → 생산배치')
run.font.size = Pt(11)
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p2 = doc.add_paragraph()
run2 = p2.add_run('※ 필수 항목은 볼드(★)로 표시되어 있습니다.\n※ 이 문서를 직접 채워서 돌려주셔도 되고, JSON 형식으로 주셔도 됩니다.')
run2.font.size = Pt(10)
run2.font.name = '맑은 고딕'

doc.add_page_break()

# ============================================================================
# 1. 거래처
# ============================================================================
add_heading_styled('1. 거래처 (공급업체 / 고객사)', level=1)
doc.add_paragraph('매입처(공급업체)와 매출처(고객사)를 등록합니다.', style='List Bullet')

add_table(
    headers=['필드명', '설명', '필수', '예시'],
    rows=[
        ['★ partnerType', '구분: supplier(공급업체) / customer(고객)', '필수', 'supplier'],
        ['★ companyName', '회사명', '필수', '(주)한국식재료'],
        ['bizNo', '사업자등록번호', '', '123-45-67890'],
        ['ceoName', '대표자명', '', '홍길동'],
        ['contactPerson', '담당자명', '', '김영업'],
        ['bizType', '업태', '', '도매'],
        ['bizItem', '종목', '', '식품원료'],
        ['address', '주소', '', '서울시 강남구...'],
        ['phone', '전화번호', '', '02-1234-5678'],
        ['email', '이메일', '', 'info@example.com'],
        ['bankName', '은행명', '', '국민은행'],
        ['bankAccount', '계좌번호', '', '123-456-789012'],
    ]
)

doc.add_paragraph()
add_heading_styled('거래처 데이터 입력란', level=2)
doc.add_paragraph('아래 표에 실제 데이터를 입력해 주세요:')

add_table(
    headers=['구분\n(supplier/customer)', '회사명', '사업자번호', '대표자', '담당자', '업태', '종목', '전화번호'],
    rows=[
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', ''],
    ]
)

doc.add_page_break()

# ============================================================================
# 2. 원재료 마스터
# ============================================================================
add_heading_styled('2. 원재료 마스터', level=1)
doc.add_paragraph('생산에 사용되는 원재료/부재료 목록을 등록합니다.', style='List Bullet')

add_table(
    headers=['필드명', '설명', '필수', '예시'],
    rows=[
        ['★ materialCode', '원재료 코드', '필수', 'RM-001'],
        ['★ materialName', '원재료명', '필수', '마카다미아'],
        ['kind', 'RAW(원재료) / MIXED(반제품)', '', 'RAW'],
        ['★ unit', '재고관리 단위', '필수', 'kg'],
        ['unitPrice', '단가 (원)', '', '25000'],
        ['purchaseUnit', '구매단위 (박스, 포 등)', '', '박스'],
        ['conversionRate', '환산비율 (1구매단위=?재고단위)', '', '10'],
        ['shelfLifeDays', '유통기한 (일)', '', '180'],
        ['safetyStockLevel', '안전재고량', '', '50'],
        ['supplierName', '주 공급업체명', '', '(주)한국식재료'],
    ]
)

doc.add_paragraph()
add_heading_styled('원재료 데이터 입력란', level=2)

add_table(
    headers=['코드', '원재료명', '구분\n(RAW/MIXED)', '재고단위', '단가(원)', '구매단위', '환산비율', '유통기한\n(일)', '안전재고', '공급업체명'],
    rows=[
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', ''],
    ]
)

doc.add_page_break()

# ============================================================================
# 3. 제품 마스터
# ============================================================================
add_heading_styled('3. 제품 마스터', level=1)
doc.add_paragraph('생산하는 완제품 목록을 등록합니다.', style='List Bullet')

add_table(
    headers=['필드명', '설명', '필수', '예시'],
    rows=[
        ['★ productCode', '제품 코드', '필수', 'FP-001'],
        ['★ productName', '제품명', '필수', '마카다미아 황참쌀떡'],
        ['category', '카테고리', '', '떡류'],
        ['unit', '단위', '', 'kg'],
        ['unitPrice', '판매단가 (원)', '', '35000'],
        ['shelfLifeDays', '유통기한 (일)', '', '14'],
    ]
)

doc.add_paragraph()
add_heading_styled('제품 데이터 입력란', level=2)

add_table(
    headers=['코드', '제품명', '카테고리', '단위', '판매단가(원)', '유통기한(일)'],
    rows=[
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
    ]
)

doc.add_page_break()

# ============================================================================
# 4. 레시피 (배합표)
# ============================================================================
add_heading_styled('4. 레시피 (배합표)', level=1)
doc.add_paragraph('제품별 원재료 배합 비율을 등록합니다. 제품 1개당 레시피 1개씩 작성하세요.', style='List Bullet')

add_table(
    headers=['필드명', '설명', '필수', '예시'],
    rows=[
        ['★ recipeName', '레시피명', '필수', '마카다미아 황참쌀떡 기본'],
        ['★ productName', '제품명 (위 제품과 매칭)', '필수', '마카다미아 황참쌀떡'],
        ['★ batchSize', '1배치 생산량', '필수', '100'],
        ['batchUnit', '단위 (기본 kg)', '', 'kg'],
        ['yieldRate', '수율 % (기본 100)', '', '95'],
        ['preparationTime', '준비시간 (분)', '', '30'],
        ['cookingTime', '조리/가공시간 (분)', '', '120'],
    ]
)

doc.add_paragraph()
add_heading_styled('레시피 헤더 입력란', level=2)

add_table(
    headers=['레시피명', '제품명', '1배치 생산량', '단위', '수율(%)', '준비시간(분)', '조리시간(분)'],
    rows=[
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
    ]
)

doc.add_paragraph()
add_heading_styled('레시피 원재료 배합 입력란 (레시피별 작성)', level=2)
doc.add_paragraph('레시피명을 적고, 해당 레시피에 들어가는 원재료를 나열하세요.')

add_table(
    headers=['레시피명', '원재료명', '투입량', '단위', '배합비율(%)'],
    rows=[
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
    ]
)

doc.add_page_break()

# ============================================================================
# 5. 원재료 입고 (매입) 내역
# ============================================================================
add_heading_styled('5. 원재료 입고 (매입) 내역', level=1)
doc.add_paragraph('원재료 매입/입고 거래 내역을 등록합니다. 회계 매입장부 + 재고 LOT가 자동 생성됩니다.', style='List Bullet')

add_table(
    headers=['필드명', '설명', '필수', '예시'],
    rows=[
        ['★ transactionDate', '거래일 (YYYY-MM-DD)', '필수', '2026-03-01'],
        ['★ supplierName', '공급업체명 (거래처와 매칭)', '필수', '(주)한국식재료'],
        ['★ itemName', '품목명 (원재료와 매칭)', '필수', '마카다미아'],
        ['★ quantity', '수량', '필수', '100'],
        ['unit', '단위', '', 'kg'],
        ['★ unitPrice', '단가 (원)', '필수', '25000'],
        ['★ totalAmount', '공급가액 (원)', '필수', '2500000'],
        ['taxAmount', '부가세 (원)', '', '250000'],
        ['lotNumber', 'LOT번호', '', 'MAC-20260301-001'],
        ['expiryDate', '유통기한 (YYYY-MM-DD)', '', '2026-08-28'],
        ['evidenceType', '증빙: tax_invoice / receipt / statement', '', 'tax_invoice'],
        ['memo', '메모', '', ''],
    ]
)

doc.add_paragraph()
add_heading_styled('매입 데이터 입력란', level=2)

add_table(
    headers=['거래일', '공급업체', '품목명', '수량', '단위', '단가', '공급가액', '부가세', 'LOT번호', '유통기한', '증빙구분'],
    rows=[
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
    ]
)

doc.add_page_break()

# ============================================================================
# 6. 생산 이력 (배치)
# ============================================================================
add_heading_styled('6. 생산 이력 (배치)', level=1)
doc.add_paragraph('생산 배치 기록을 등록합니다. 자동으로 일일일지 및 재고가 연동됩니다.', style='List Bullet')

add_table(
    headers=['필드명', '설명', '필수', '예시'],
    rows=[
        ['batchCode', '배치코드 (자동생성 가능)', '', 'B-20260311-001'],
        ['★ productName', '제품명 (제품 마스터와 매칭)', '필수', '마카다미아 황참쌀떡'],
        ['recipeName', '레시피명', '', '마카다미아 황참쌀떡 기본'],
        ['★ plannedDate', '생산일 (YYYY-MM-DD)', '필수', '2026-03-11'],
        ['★ plannedQuantity', '계획수량', '필수', '300'],
        ['actualQuantity', '실 생산량', '', '285'],
        ['status', 'planned / completed / in_progress', '', 'completed'],
        ['startTime', '시작시간 (HH:mm)', '', '08:00'],
        ['endTime', '종료시간 (HH:mm)', '', '14:00'],
        ['lotNumber', '완제품 LOT번호', '', 'FP-20260311-001'],
        ['expiryDate', '유통기한 (YYYY-MM-DD)', '', '2026-03-25'],
        ['notes', '비고', '', ''],
    ]
)

doc.add_paragraph()
add_heading_styled('생산 배치 데이터 입력란', level=2)

add_table(
    headers=['제품명', '레시피명', '생산일', '계획수량', '실생산량', '상태', '시작', '종료', 'LOT번호', '유통기한', '비고'],
    rows=[
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
        ['', '', '', '', '', '', '', '', '', '', ''],
    ]
)

doc.add_paragraph()
add_heading_styled('배치별 원재료 사용 내역 (선택)', level=2)
doc.add_paragraph('각 배치에서 어떤 원재료를 얼마나 사용했는지 기록합니다. (비워도 레시피 기반 자동 계산 가능)')

add_table(
    headers=['배치 (제품명+생산일)', '원재료명', '사용 LOT번호', '사용량', '단위'],
    rows=[
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
    ]
)

doc.add_page_break()

# ============================================================================
# 참고: JSON 형식 예시
# ============================================================================
add_heading_styled('참고: JSON 형식 예시', level=1)
doc.add_paragraph('위 표 대신 아래와 같은 JSON 형식으로 데이터를 주셔도 됩니다.')

add_json_block('''{
  "partners": [
    {
      "partnerType": "supplier",
      "companyName": "(주)한국식재료",
      "bizNo": "123-45-67890",
      "ceoName": "홍길동",
      "phone": "02-1234-5678"
    }
  ],
  "materials": [
    {
      "materialCode": "RM-001",
      "materialName": "마카다미아",
      "unit": "kg",
      "unitPrice": 25000,
      "shelfLifeDays": 180
    }
  ],
  "products": [
    {
      "productCode": "FP-001",
      "productName": "마카다미아 황참쌀떡",
      "unit": "kg",
      "unitPrice": 35000,
      "shelfLifeDays": 14
    }
  ],
  "recipes": [
    {
      "recipeName": "마카다미아 황참쌀떡 기본",
      "productName": "마카다미아 황참쌀떡",
      "batchSize": 100,
      "lines": [
        { "materialName": "마카다미아", "quantity": 30, "unit": "kg" }
      ]
    }
  ],
  "purchases": [
    {
      "transactionDate": "2026-03-01",
      "supplierName": "(주)한국식재료",
      "itemName": "마카다미아",
      "quantity": 100,
      "unitPrice": 25000,
      "totalAmount": 2500000,
      "taxAmount": 250000
    }
  ],
  "batches": [
    {
      "productName": "마카다미아 황참쌀떡",
      "plannedDate": "2026-03-11",
      "plannedQuantity": 300,
      "actualQuantity": 285,
      "status": "completed"
    }
  ]
}''')

doc.add_paragraph()
add_note('입력 순서: partners → materials → products → recipes → purchases → batches (FK 의존성)')
add_note('이름으로 매칭합니다: supplierName ↔ companyName, productName ↔ productName, materialName ↔ materialName')
add_note('LOT번호를 비워두면 자동 생성됩니다.')
add_note('이 워드 문서의 표를 직접 채워 주셔도 되고, JSON 파일로 주셔도 됩니다.')

# ── 저장 ──
output_path = '/home/user/haccp_v3/HACCP-ONE_데이터입력_템플릿.docx'
doc.save(output_path)
print(f'✅ 워드 문서 생성 완료: {output_path}')
